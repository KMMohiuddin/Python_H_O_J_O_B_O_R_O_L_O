import os
import docx
import datetime
from datetime import datetime

def delete_row_in_table(document, search_text):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    table._tbl.remove(row._tr)
                    return True
    return False

def merge_docx_files(path):
    docx_files = [file for file in os.listdir(path) if file.endswith('.docx')]
    merged_doc = docx.Document()
    
    report_template = docx.Document(report_temp)
    report_date = datetime.now().strftime("%dth %B %Y")

    for table in report_template.tables:
        for row in table.rows:
            for cell in row.cells:
                #print("hello")
                if "DATE:" in cell.text:
                    #print(cell.text)
                    cell.text = cell.text.replace(cell.text, f"Date: {report_date}")
                #print(cell.text)
                merged_table = merged_doc.add_table(rows=0, cols=0)
                #merged_table._cells.append(table._cells)
                merged_table._element.append(table._element)


    for file in docx_files:
        doc = docx.Document(os.path.join(path, file))
        for table in doc.tables:
            for row in table.rows:
                if delete_row_in_table(doc, search_text):
                    merged_table = merged_doc.add_table(rows=0, cols=0)
                    merged_table._element.append(table._element)



    return merged_doc

directory_list = ["business_intelligence","business_performance","customer_service","internal_audit","process_development","reconciliation","revenue_assurance"]
op_directory = "combined_reports"
report_temp = "report_template.docx"
search_text = 'Date:'
current_directory = os.getcwd()


merged_document = merge_docx_files(current_directory)
current_date = datetime.now().strftime('%Y%m%d')
merged_filename = f'{current_date}.docx'
merged_document.save(os.path.join(current_directory, merged_filename))
